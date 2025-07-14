from docx import Document
from docx.shared import RGBColor


def export_text_with_highlights(text, keywords, output_path):
    """
    将匹配到关键词的文本导出为 Word 文档，并高亮关键词（字体颜色变红）。

    :param text: str, 原始文本内容
    :param keywords: list[str], 要高亮的关键词
    :param output_path: str, Word文档保存路径
    """
    doc = Document()
    para = doc.add_paragraph()

    # 优先匹配长关键词
    keywords = sorted(set(filter(None, keywords)), key=len, reverse=True)

    i = 0
    while i < len(text):
        matched = False
        for keyword in keywords:
            if text[i:i + len(keyword)] == keyword:
                run = para.add_run(keyword)
                run.font.color.rgb = RGBColor(255, 0, 0)  # 红色高亮
                matched = True
                i += len(keyword)
                break
        if not matched:
            para.add_run(text[i])
            i += 1

    doc.save(output_path)
